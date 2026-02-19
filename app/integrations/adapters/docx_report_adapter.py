"""
DOCX report adapter using python-docx.
"""

from __future__ import annotations

import io
import os
import tempfile
from typing import Any, Dict, List, Optional

from ..base import ToolAdapter, ToolContext, ToolParameter, ToolResult, ToolSpec
from ..plotting import apply_scienceplots_style


class DocxReportAdapter(ToolAdapter):
    spec = ToolSpec(
        id="docx_report",
        name="DOCX Report",
        version="1.0",
        description="Generate a DOCX report with figures and tables.",
        input_types=["image", "csv", "txt"],
        parameters=[
            ToolParameter("title", "string", False, "Analysis Report", "Report title"),
            ToolParameter("summary", "string", False, None, "Short summary paragraph"),
            ToolParameter("template_path", "string", False, None, "DOCX template path"),
            ToolParameter("table_rows", "array", False, None, "List of table row dicts"),
            ToolParameter("table_columns", "array", False, None, "Explicit table column order"),
            ToolParameter("plot_x_col", "string", False, None, "CSV X column name"),
            ToolParameter("plot_y_col", "string", False, None, "CSV Y column name"),
            ToolParameter("report_name", "string", False, None, "Report filename stem"),
            ToolParameter("image_caption", "string", False, None, "Caption for embedded image"),
        ],
        outputs=["report_path", "image_count", "table_rows"],
    )

    def run(self, context: ToolContext) -> ToolResult:
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError as exc:
            return ToolResult(status="failed", error=f"python-docx not available: {exc}")

        params = context.parameters or {}
        title = params.get("title", "Analysis Report")
        summary = params.get("summary")
        template_path = params.get("template_path")
        report_name = params.get("report_name")
        image_caption = params.get("image_caption")

        if template_path:
            if not os.path.exists(template_path):
                return ToolResult(status="failed", error="template_path not found")
            document = Document(template_path)
        else:
            document = Document()

        document.add_heading(str(title), level=1)
        if summary:
            document.add_paragraph(str(summary))

        image_count = 0
        warnings: List[str] = []

        if context.file_bytes:
            image_path = self._prepare_image(context, params, warnings)
            if image_path:
                try:
                    document.add_picture(image_path, width=Inches(5.5))
                    image_count += 1
                    if image_caption:
                        document.add_paragraph(str(image_caption))
                finally:
                    if image_path and os.path.exists(image_path):
                        os.remove(image_path)

        table_rows = params.get("table_rows") or []
        if not isinstance(table_rows, list):
            warnings.append("table_rows must be a list")
            table_rows = []
        table_columns = params.get("table_columns")
        if table_rows:
            columns = self._resolve_columns(table_rows, table_columns)
            table = document.add_table(rows=1, cols=len(columns))
            table.style = "Light Grid"
            header_cells = table.rows[0].cells
            for idx, col in enumerate(columns):
                header_cells[idx].text = str(col)
            for row in table_rows:
                cells = table.add_row().cells
                for idx, col in enumerate(columns):
                    cells[idx].text = str(row.get(col, ""))

        output_dir = os.getenv("REPORT_PATH", "data/outputs/reports")
        os.makedirs(output_dir, exist_ok=True)
        if not report_name:
            stem = os.path.splitext(context.filename or "report")[0]
            report_name = f"{stem}_report"
        report_path = os.path.join(output_dir, f"{report_name}.docx")

        try:
            document.save(report_path)
        except Exception as exc:
            return ToolResult(status="failed", error=f"Failed to save DOCX: {exc}")

        output = {
            "report_path": report_path,
            "image_count": image_count,
            "table_rows": len(table_rows),
        }
        annotations = [
            {
                "analysis": "docx_report",
                "report_path": report_path,
                "image_count": image_count,
                "table_rows": len(table_rows),
            }
        ]
        conclusion = "DOCX report generated"

        return ToolResult(
            status="completed",
            output=output,
            annotations=annotations,
            conclusion=conclusion,
            warnings=warnings,
        )

    def _prepare_image(
        self, context: ToolContext, params: Dict[str, Any], warnings: List[str]
    ) -> Optional[str]:
        filename = (context.filename or "").lower()
        if filename.endswith((".png", ".jpg", ".jpeg", ".tif", ".tiff")):
            return self._write_temp_file(context.file_bytes, suffix=os.path.splitext(filename)[1])

        try:
            import pandas as pd
            import matplotlib.pyplot as plt
        except ImportError as exc:
            warnings.append(f"Plot skipped (missing dependency): {exc}")
            return None

        try:
            text = context.file_bytes.decode("utf-8", errors="replace")
            df = pd.read_csv(io.StringIO(text), sep=None, engine="python")
        except Exception as exc:
            warnings.append(f"Plot skipped (CSV parse failed): {exc}")
            return None

        if df.empty or len(df.columns) < 2:
            warnings.append("Plot skipped (not enough columns)")
            return None

        x_col = params.get("plot_x_col")
        y_col = params.get("plot_y_col")
        col_map = {c.lower(): c for c in df.columns}
        if x_col:
            x_col = col_map.get(str(x_col).lower(), x_col)
        else:
            x_col = df.columns[0]
        if y_col:
            y_col = col_map.get(str(y_col).lower(), y_col)
        else:
            y_col = df.columns[1]

        if x_col not in df.columns or y_col not in df.columns:
            warnings.append("Plot skipped (columns not found)")
            return None

        applied, error = apply_scienceplots_style()
        if not applied and error:
            warnings.append(error)

        plt.figure(figsize=(6, 4))
        plt.plot(df[x_col], df[y_col], linewidth=1.5)
        plt.xlabel(str(x_col))
        plt.ylabel(str(y_col))
        plt.tight_layout()

        plot_path = self._write_temp_file(b"", suffix=".png")
        plt.savefig(plot_path, dpi=200)
        plt.close()
        return plot_path

    def _resolve_columns(self, rows: List[Dict[str, Any]], columns: Optional[List[str]]) -> List[str]:
        if columns:
            return list(columns)
        return list(rows[0].keys())

    def _write_temp_file(self, data: bytes, suffix: str = "") -> str:
        handle = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        try:
            if data:
                handle.write(data)
            return handle.name
        finally:
            handle.close()
